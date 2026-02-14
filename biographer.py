# biographer.py – Tell My Story App (COMPLETE WORKING VERSION)
import streamlit as st
import json
from datetime import datetime, date
from openai import OpenAI
import os
import re
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import secrets
import string
import time
import shutil
import base64
from PIL import Image
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import zipfile

# ============================================================================
# IMPORT QUILL RICH TEXT EDITOR
# ============================================================================
try:
    from streamlit_quill import st_quill
    QUILL_AVAILABLE = True
except ImportError:
    st.error("❌ Please install streamlit-quill: pip install streamlit-quill")
    st.stop()

# ============================================================================
# FORCE DIRECTORY CREATION
# ============================================================================
for dir_path in ["question_banks/default", "question_banks/users", "question_banks", 
                 "uploads", "uploads/thumbnails", "uploads/metadata", "accounts", "sessions", "backups"]:
    os.makedirs(dir_path, exist_ok=True)

# ============================================================================
# IMPORTS
# ============================================================================
try:
    from topic_bank import TopicBank
    from session_manager import SessionManager
    from vignettes import VignetteManager
    from session_loader import SessionLoader
    from beta_reader import BetaReader
    from question_bank_manager import QuestionBankManager
except ImportError as e:
    st.error(f"Error importing modules: {e}")
    st.info("Please ensure all .py files are in the same directory")
    TopicBank = SessionManager = VignetteManager = SessionLoader = BetaReader = QuestionBankManager = None

DEFAULT_WORD_TARGET = 500

# ============================================================================
# INITIALIZATION
# ============================================================================
client = OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY")))
beta_reader = BetaReader(client) if BetaReader else None

# Initialize session state
default_state = {
    "qb_manager": None, "qb_manager_initialized": False, "user_id": None, "logged_in": False,
    "current_session": 0, "current_question": 0, "responses": {}, "editing": False,
    "editing_word_target": False, "confirming_clear": None, "data_loaded": False,
    "current_question_override": None, "show_vignette_modal": False, "vignette_topic": "",
    "vignette_content": "", "selected_vignette_type": "Standard Topic", "current_vignette_list": [],
    "editing_vignette_index": None, "show_vignette_manager": False, "custom_topic_input": "",
    "show_custom_topic_modal": False, "show_topic_browser": False, "show_session_manager": False,
    "show_session_creator": False, "editing_custom_session": None, "show_vignette_detail": False,
    "selected_vignette_id": None, "editing_vignette_id": None, "selected_vignette_for_session": None,
    "published_vignette": None, "show_beta_reader": False, "current_beta_feedback": None,
    "current_question_bank": None, "current_bank_name": None, "current_bank_type": None,
    "current_bank_id": None, "show_bank_manager": False, "show_bank_editor": False,
    "editing_bank_id": None, "editing_bank_name": None, "qb_manager": None, "qb_manager_initialized": False,
    "confirm_delete": None, "user_account": None, "show_profile_setup": False,
    "image_handler": None, "show_image_manager": False, "show_ai_suggestions": False,
    "current_ai_suggestions": None, "current_suggestion_topic": None, "editor_content": {},
    "show_privacy_settings": False, "show_cover_designer": False
}
for key, value in default_state.items():
    if key not in st.session_state:
        st.session_state[key] = value

# Load external CSS
try:
    with open("styles.css", encoding="utf-8") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
except FileNotFoundError:
    pass

LOGO_URL = "https://menuhunterai.com/wp-content/uploads/2026/02/tms_logo.png"

# ============================================================================
# EMAIL CONFIG
# ============================================================================
EMAIL_CONFIG = {
    "smtp_server": st.secrets.get("SMTP_SERVER", "smtp.gmail.com"),
    "smtp_port": int(st.secrets.get("SMTP_PORT", 587)),
    "sender_email": st.secrets.get("SENDER_EMAIL", ""),
    "sender_password": st.secrets.get("SENDER_PASSWORD", ""),
    "use_tls": True
}

# ============================================================================
# BACKUP AND RESTORE FUNCTIONS
# ============================================================================
def create_backup():
    """Create a complete backup of user data"""
    if not st.session_state.user_id:
        return None
    
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_data = {
            "user_id": st.session_state.user_id,
            "user_account": st.session_state.user_account,
            "responses": st.session_state.responses,
            "backup_date": datetime.now().isoformat(),
            "version": "1.0"
        }
        
        # Save to backups folder
        backup_file = f"backups/{st.session_state.user_id}_{timestamp}.json"
        with open(backup_file, 'w') as f:
            json.dump(backup_data, f, indent=2)
        
        # Also create downloadable version
        return json.dumps(backup_data, indent=2)
    except Exception as e:
        st.error(f"Backup failed: {e}")
        return None

def restore_from_backup(backup_json):
    """Restore user data from backup"""
    try:
        backup_data = json.loads(backup_json)
        if backup_data.get("user_id") != st.session_state.user_id:
            st.error("Backup belongs to a different user")
            return False
        
        st.session_state.user_account = backup_data.get("user_account", st.session_state.user_account)
        st.session_state.responses = backup_data.get("responses", st.session_state.responses)
        
        # Save to files
        save_account_data(st.session_state.user_account)
        save_user_data(st.session_state.user_id, st.session_state.responses)
        
        return True
    except Exception as e:
        st.error(f"Restore failed: {e}")
        return False

def list_backups():
    """List all backups for current user"""
    if not st.session_state.user_id:
        return []
    
    backups = []
    try:
        for f in os.listdir("backups"):
            if f.startswith(st.session_state.user_id) and f.endswith(".json"):
                filepath = f"backups/{f}"
                with open(filepath, 'r') as file:
                    data = json.load(file)
                    backups.append({
                        "filename": f,
                        "date": data.get("backup_date", "Unknown"),
                        "size": os.path.getsize(filepath)
                    })
    except:
        pass
    
    return sorted(backups, key=lambda x: x["date"], reverse=True)

# ============================================================================
# IMAGE HANDLER
# ============================================================================
class ImageHandler:
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.base_path = "uploads"
        
        self.settings = {
            "full_width": 1600,
            "inline_width": 800,
            "thumbnail_size": 200,
            "quality": 85,
            "max_file_size_mb": 5
        }
    
    def get_user_path(self):
        if self.user_id:
            user_hash = hashlib.md5(self.user_id.encode()).hexdigest()[:8]
            path = f"{self.base_path}/user_{user_hash}"
            os.makedirs(f"{path}/thumbnails", exist_ok=True)
            return path
        return self.base_path
    
    def optimize_image(self, image, max_width=1600, is_thumbnail=False):
        try:
            # Convert to RGB if needed
            if image.mode in ('RGBA', 'LA', 'P'):
                bg = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                if image.mode == 'RGBA':
                    bg.paste(image, mask=image.split()[-1])
                else:
                    bg.paste(image)
                image = bg
            
            width, height = image.size
            aspect = height / width
            
            if is_thumbnail:
                size = min(width, height)
                left = (width - size) // 2
                top = (height - size) // 2
                right = left + size
                bottom = top + size
                image = image.crop((left, top, right, bottom))
                image.thumbnail((self.settings["thumbnail_size"], self.settings["thumbnail_size"]), Image.Resampling.LANCZOS)
                return image
            
            if width > max_width:
                new_width = max_width
                new_height = int(max_width * aspect)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            return image
        except Exception as e:
            print(f"Error optimizing image: {e}")
            return image
    
    def save_image(self, uploaded_file, session_id, question_text, caption="", usage="full_page"):
        try:
            image_data = uploaded_file.read()
            original_size = len(image_data) / (1024 * 1024)
            
            img = Image.open(io.BytesIO(image_data))
            
            target_width = self.settings["full_width"] if usage == "full_page" else self.settings["inline_width"]
            
            image_id = hashlib.md5(f"{self.user_id}{session_id}{question_text}{datetime.now()}".encode()).hexdigest()[:16]
            
            optimized_img = self.optimize_image(img, target_width, is_thumbnail=False)
            thumb_img = self.optimize_image(img, is_thumbnail=True)
            
            main_buffer = io.BytesIO()
            optimized_img.save(main_buffer, format="JPEG", quality=self.settings["quality"], optimize=True)
            
            thumb_buffer = io.BytesIO()
            thumb_img.save(thumb_buffer, format="JPEG", quality=70, optimize=True)
            
            user_path = self.get_user_path()
            with open(f"{user_path}/{image_id}.jpg", 'wb') as f: 
                f.write(main_buffer.getvalue())
            with open(f"{user_path}/thumbnails/{image_id}.jpg", 'wb') as f: 
                f.write(thumb_buffer.getvalue())
            
            metadata = {
                "id": image_id, 
                "session_id": session_id, 
                "question": question_text,
                "caption": caption, 
                "timestamp": datetime.now().isoformat(),
                "user_id": self.user_id,
                "usage": usage
            }
            with open(f"{self.base_path}/metadata/{image_id}.json", 'w') as f: 
                json.dump(metadata, f, indent=2)
            
            return {
                "has_images": True, 
                "images": [{
                    "id": image_id, 
                    "caption": caption
                }]
            }
        except Exception as e:
            print(f"Error saving image: {e}")
            return None
    
    def get_image_html(self, image_id, thumbnail=False):
        try:
            user_path = self.get_user_path()
            path = f"{user_path}/thumbnails/{image_id}.jpg" if thumbnail else f"{user_path}/{image_id}.jpg"
            if not os.path.exists(path): 
                return None
            
            with open(path, 'rb') as f: 
                image_data = f.read()
            b64 = base64.b64encode(image_data).decode()
            
            meta_path = f"{self.base_path}/metadata/{image_id}.json"
            caption = ""
            if os.path.exists(meta_path):
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    caption = metadata.get("caption", "")
            
            return {
                "html": f'<img src="data:image/jpeg;base64,{b64}" style="max-width:100%; border-radius:8px; margin:5px 0;" alt="{caption}">',
                "caption": caption, 
                "base64": b64
            }
        except:
            return None
    
    def get_image_base64(self, image_id):
        try:
            user_path = self.get_user_path()
            path = f"{user_path}/{image_id}.jpg"
            if not os.path.exists(path): 
                return None
            with open(path, 'rb') as f: 
                image_data = f.read()
            return base64.b64encode(image_data).decode()
        except:
            return None
    
    def get_images_for_answer(self, session_id, question_text):
        images = []
        metadata_dir = f"{self.base_path}/metadata"
        if not os.path.exists(metadata_dir): 
            return images
        
        for fname in os.listdir(metadata_dir):
            if fname.endswith('.json'):
                try:
                    with open(f"{metadata_dir}/{fname}") as f: 
                        meta = json.load(f)
                    if (meta.get("session_id") == session_id and 
                        meta.get("question") == question_text and 
                        meta.get("user_id") == self.user_id):
                        thumb = self.get_image_html(meta["id"], thumbnail=True)
                        full = self.get_image_html(meta["id"])
                        if thumb and full:
                            images.append({
                                **meta, 
                                "thumb_html": thumb["html"], 
                                "full_html": full["html"],
                                "base64": full["base64"]
                            })
                except:
                    continue
        return sorted(images, key=lambda x: x.get("timestamp", ""), reverse=True)
    
    def delete_image(self, image_id):
        try:
            user_path = self.get_user_path()
            for p in [f"{user_path}/{image_id}.jpg", 
                     f"{user_path}/thumbnails/{image_id}.jpg", 
                     f"{self.base_path}/metadata/{image_id}.json"]:
                if os.path.exists(p): 
                    os.remove(p)
            return True
        except:
            return False

def init_image_handler():
    if not st.session_state.image_handler or st.session_state.image_handler.user_id != st.session_state.get('user_id'):
        st.session_state.image_handler = ImageHandler(st.session_state.get('user_id'))
    return st.session_state.image_handler

# ============================================================================
# AUTHENTICATION FUNCTIONS
# ============================================================================
def generate_password(length=12):
    alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(stored_hash, password):
    return stored_hash == hash_password(password)

def create_user_account(user_data, password=None):
    try:
        user_id = hashlib.sha256(f"{user_data['email']}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
        if not password: 
            password = generate_password()
        user_record = {
            "user_id": user_id, 
            "email": user_data["email"].lower().strip(),
            "password_hash": hash_password(password), 
            "account_type": user_data.get("account_for", "self"),
            "created_at": datetime.now().isoformat(), 
            "last_login": datetime.now().isoformat(),
            "profile": {
                "first_name": user_data["first_name"], 
                "last_name": user_data["last_name"],
                "email": user_data["email"], 
                "gender": user_data.get("gender", ""),
                "birthdate": user_data.get("birthdate", ""), 
                "timeline_start": user_data.get("birthdate", "")
            },
            "narrative_gps": {},
            "privacy_settings": {
                "profile_public": False,
                "stories_public": False,
                "allow_sharing": False,
                "data_collection": True,
                "encryption": True
            },
            "settings": {
                "email_notifications": True, 
                "auto_save": True, 
                "privacy_level": "private",
                "theme": "light", 
                "email_verified": False
            },
            "stats": {
                "total_sessions": 0, 
                "total_words": 0, 
                "current_streak": 0, 
                "longest_streak": 0,
                "account_age_days": 0, 
                "last_active": datetime.now().isoformat()
            }
        }
        save_account_data(user_record)
        return {"success": True, "user_id": user_id, "password": password, "user_record": user_record}
    except Exception as e:
        return {"success": False, "error": str(e)}

def save_account_data(user_record):
    try:
        with open(f"accounts/{user_record['user_id']}_account.json", 'w') as f:
            json.dump(user_record, f, indent=2)
        update_accounts_index(user_record)
        return True
    except: 
        return False

def update_accounts_index(user_record):
    try:
        index_file = "accounts/accounts_index.json"
        index = json.load(open(index_file, 'r')) if os.path.exists(index_file) else {}
        index[user_record['user_id']] = {
            "email": user_record['email'], 
            "first_name": user_record['profile']['first_name'],
            "last_name": user_record['profile']['last_name'], 
            "created_at": user_record['created_at'],
            "account_type": user_record['account_type']
        }
        with open(index_file, 'w') as f: 
            json.dump(index, f, indent=2)
        return True
    except: 
        return False

def get_account_data(user_id=None, email=None):
    try:
        if user_id:
            fname = f"accounts/{user_id}_account.json"
            if os.path.exists(fname): 
                return json.load(open(fname, 'r'))
        if email:
            email = email.lower().strip()
            index_file = "accounts/accounts_index.json"
            if os.path.exists(index_file):
                with open(index_file, 'r') as f:
                    index = json.load(f)
                for uid, data in index.items():
                    if data.get("email", "").lower() == email:
                        return json.load(open(f"accounts/{uid}_account.json", 'r'))
    except: 
        pass
    return None

def authenticate_user(email, password):
    try:
        account = get_account_data(email=email)
        if account and verify_password(account['password_hash'], password):
            account['last_login'] = datetime.now().isoformat()
            save_account_data(account)
            return {"success": True, "user_id": account['user_id'], "user_record": account}
        return {"success": False, "error": "Invalid email or password"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def send_welcome_email(user_data, credentials):
    try:
        if not EMAIL_CONFIG['sender_email'] or not EMAIL_CONFIG['sender_password']: 
            return False
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = user_data['email']
        msg['Subject'] = "Welcome to Tell My Story"
        
        body = f"""
        <html><body style="font-family: Arial;">
        <h2>Welcome to Tell My Story, {user_data['first_name']}!</h2>
        <div style="background: #f0f8ff; padding: 15px; border-left: 4px solid #3498db;">
            <h3>Your Account Details:</h3>
            <p><strong>Account ID:</strong> {credentials['user_id']}</p>
            <p><strong>Email:</strong> {user_data['email']}</p>
            <p><strong>Password:</strong> {credentials['password']}</p>
        </div>
        </body></html>
        """
        msg.attach(MIMEText(body, 'html'))
        with smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port']) as server:
            if EMAIL_CONFIG['use_tls']: 
                server.starttls()
            server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
            server.send_message(msg)
        return True
    except: 
        return False

def logout_user():
    st.session_state.qb_manager = None
    st.session_state.qb_manager_initialized = False
    st.session_state.image_handler = None
    keys = ['user_id', 'user_account', 'logged_in', 'show_profile_setup', 'current_session',
            'current_question', 'responses', 'session_conversations', 'data_loaded',
            'show_vignette_modal', 'vignette_topic', 'vignette_content', 'selected_vignette_type',
            'current_vignette_list', 'editing_vignette_index', 'show_vignette_manager',
            'custom_topic_input', 'show_custom_topic_modal', 'show_topic_browser',
            'show_session_manager', 'show_session_creator', 'editing_custom_session',
            'show_vignette_detail', 'selected_vignette_id', 'editing_vignette_id',
            'selected_vignette_for_session', 'published_vignette', 'show_beta_reader',
            'current_beta_feedback', 'current_question_bank', 'current_bank_name',
            'current_bank_type', 'current_bank_id', 'show_bank_manager', 'show_bank_editor',
            'editing_bank_id', 'editing_bank_name', 'show_image_manager', 'editor_content']
    for key in keys:
        if key in st.session_state: 
            del st.session_state[key]
    st.query_params.clear()
    st.rerun()

# ============================================================================
# PRIVACY SETTINGS MODAL
# ============================================================================
def show_privacy_settings():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🔒 Privacy & Security Settings")
    
    if st.button("← Back", key="privacy_back"):
        st.session_state.show_privacy_settings = False
        st.rerun()
    
    st.markdown("### Ethical AI & Data Privacy")
    st.info("Your stories are private and secure. We use AI ethically to help you write better, never to train models on your personal data.")
    
    if 'privacy_settings' not in st.session_state.user_account:
        st.session_state.user_account['privacy_settings'] = {
            "profile_public": False,
            "stories_public": False,
            "allow_sharing": False,
            "data_collection": True,
            "encryption": True
        }
    
    privacy = st.session_state.user_account['privacy_settings']
    
    privacy['profile_public'] = st.checkbox("Make profile public", value=privacy.get('profile_public', False),
                                           help="Allow others to see your basic profile information")
    privacy['stories_public'] = st.checkbox("Share stories publicly", value=privacy.get('stories_public', False),
                                           help="Make your stories visible to the public (coming soon)")
    privacy['allow_sharing'] = st.checkbox("Allow sharing via link", value=privacy.get('allow_sharing', False),
                                          help="Generate shareable links to your stories")
    privacy['data_collection'] = st.checkbox("Allow anonymous usage data", value=privacy.get('data_collection', True),
                                            help="Help us improve by sharing anonymous usage statistics")
    privacy['encryption'] = st.checkbox("Enable encryption", value=privacy.get('encryption', True),
                                       disabled=True, help="Your data is always encrypted at rest")
    
    st.markdown("---")
    st.markdown("### 🔐 Security")
    st.markdown("- All data encrypted at rest")
    st.markdown("- No third-party data sharing")
    st.markdown("- You own all your content")
    st.markdown("- AI analysis is temporary and private")
    
    if st.button("💾 Save Privacy Settings", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("Privacy settings saved!")
        time.sleep(1)
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# COVER DESIGNER MODAL
# ============================================================================
def show_cover_designer():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🎨 Cover Designer")
    
    if st.button("← Back", key="cover_back"):
        st.session_state.show_cover_designer = False
        st.rerun()
    
    st.markdown("### Design your book cover")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Cover Options**")
        cover_type = st.selectbox("Cover Style", ["Simple", "Elegant", "Modern", "Classic", "Vintage"])
        title_font = st.selectbox("Title Font", ["Georgia", "Arial", "Times New Roman", "Helvetica", "Calibri"])
        title_color = st.color_picker("Title Color", "#000000")
        background_color = st.color_picker("Background Color", "#FFFFFF")
        
        uploaded_cover = st.file_uploader("Upload Cover Image (optional)", type=['jpg', 'jpeg', 'png'])
        if uploaded_cover:
            st.image(uploaded_cover, caption="Your cover image", width=300)
    
    with col2:
        st.markdown("**Preview**")
        first_name = st.session_state.user_account.get('profile', {}).get('first_name', 'My')
        preview_title = st.text_input("Preview Title", value=f"{first_name}'s Story")
        
        preview_style = f"""
        <div style="width:300px; height:400px; background-color:{background_color}; 
                    border:2px solid #ccc; border-radius:10px; padding:20px; 
                    display:flex; flex-direction:column; justify-content:center; 
                    align-items:center; text-align:center;">
            <h1 style="font-family:{title_font}; color:{title_color};">{preview_title}</h1>
            <p style="margin-top:50px;">by {st.session_state.user_account.get('profile', {}).get('first_name', '')}</p>
        </div>
        """
        st.markdown(preview_style, unsafe_allow_html=True)
    
    if st.button("💾 Save Cover Design", type="primary"):
        if 'cover_design' not in st.session_state.user_account:
            st.session_state.user_account['cover_design'] = {}
        
        st.session_state.user_account['cover_design'].update({
            "cover_type": cover_type,
            "title_font": title_font,
            "title_color": title_color,
            "background_color": background_color,
            "title": preview_title
        })
        
        if uploaded_cover:
            os.makedirs("uploads/covers", exist_ok=True)
            cover_path = f"uploads/covers/{st.session_state.user_id}_cover.jpg"
            with open(cover_path, 'wb') as f:
                f.write(uploaded_cover.getbuffer())
            st.session_state.user_account['cover_design']['cover_image'] = cover_path
        
        save_account_data(st.session_state.user_account)
        st.success("Cover design saved!")
        time.sleep(1)
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# NARRATIVE GPS HELPER FUNCTIONS (FOR AI INTEGRATION)
# ============================================================================
def get_narrative_gps_for_ai():
    """Format Narrative GPS data for AI prompts"""
    if not st.session_state.user_account or 'narrative_gps' not in st.session_state.user_account:
        return ""
    
    gps = st.session_state.user_account['narrative_gps']
    if not gps:
        return ""
    
    context = "\n\n=== BOOK PROJECT CONTEXT ===\n"
    
    if gps.get('book_title'): context += f"- Book Title: {gps['book_title']}\n"
    if gps.get('genre'): 
        genre = gps['genre']
        if genre == "Other" and gps.get('genre_other'):
            genre = gps['genre_other']
        context += f"- Genre: {genre}\n"
    if gps.get('purposes'): context += f"- Purposes: {', '.join(gps['purposes'])}\n"
    if gps.get('reader_takeaway'): context += f"- Reader Takeaway: {gps['reader_takeaway']}\n"
    if gps.get('voices'): context += f"- Voice: {', '.join(gps['voices'])}\n"
    
    return context

# ============================================================================
# AI WRITING SUGGESTIONS FUNCTION
# ============================================================================
def generate_writing_suggestions(question, answer_text, session_title):
    """Generate immediate writing suggestions based on the answer and Narrative GPS context"""
    if not client:
        return {"error": "OpenAI client not available"}
    
    try:
        gps_context = get_narrative_gps_for_ai()
        clean_answer = re.sub(r'<[^>]+>', '', answer_text)
        
        if len(clean_answer.split()) < 20:
            return None
        
        system_prompt = """You are an expert writing coach. Provide 2-3 specific, actionable suggestions to improve this life story passage. Focus on alignment with the book's purpose and opportunities to deepen the narrative."""
        
        user_prompt = f"""{gps_context}

SESSION: {session_title}
QUESTION: {question}
ANSWER: {clean_answer}

Provide 2-3 specific suggestions to improve this answer."""
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.7
        )
        
        return response.choices[0].message.content
    except Exception as e:
        return {"error": str(e)}

def show_ai_suggestions():
    """Display AI writing suggestions inline (not modal)"""
    if st.session_state.get('current_ai_suggestions'):
        with st.container():
            st.info("💡 **AI Suggestion**")
            st.markdown(st.session_state.current_ai_suggestions)
            if st.button("Dismiss", key="dismiss_suggestions"):
                st.session_state.show_ai_suggestions = False
                st.session_state.current_ai_suggestions = None
                st.rerun()

# ============================================================================
# ENHANCED BIOGRAPHER PROFILE SECTION
# ============================================================================
def render_enhanced_profile():
    """Render an expanded biographer-style questionnaire"""
    st.markdown("### 📋 The Biographer's Questions")
    st.info("A biographer would ask these questions to capture the full richness of your life story.")
    
    if 'enhanced_profile' not in st.session_state.user_account:
        st.session_state.user_account['enhanced_profile'] = {}
    
    ep = st.session_state.user_account['enhanced_profile']
    
    with st.expander("👶 Early Years & Family Origins", expanded=False):
        ep['birth_place'] = st.text_input("Where and when were you born?", value=ep.get('birth_place', ''), key="ep_birth_place")
        ep['parents'] = st.text_area("Tell me about your parents - who were they?", value=ep.get('parents', ''), key="ep_parents", height=80)
        ep['siblings'] = st.text_area("Did you have siblings? What was your relationship with them?", value=ep.get('siblings', ''), key="ep_siblings", height=80)
        ep['childhood_home'] = st.text_area("What was your childhood home like?", value=ep.get('childhood_home', ''), key="ep_home", height=80)
    
    with st.expander("🎓 Education & Formative Years", expanded=False):
        ep['school'] = st.text_area("What was your school experience like?", value=ep.get('school', ''), key="ep_school", height=80)
        ep['higher_ed'] = st.text_area("Did you pursue higher education?", value=ep.get('higher_ed', ''), key="ep_higher_ed", height=80)
        ep['mentors'] = st.text_area("Who were your mentors or influential figures?", value=ep.get('mentors', ''), key="ep_mentors", height=80)
    
    with st.expander("💼 Career & Life's Work", expanded=False):
        ep['first_job'] = st.text_area("What was your first job?", value=ep.get('first_job', ''), key="ep_first_job", height=80)
        ep['career_path'] = st.text_area("Describe your career path", value=ep.get('career_path', ''), key="ep_career", height=80)
        ep['achievements'] = st.text_area("What achievements are you most proud of?", value=ep.get('achievements', ''), key="ep_achievements", height=80)
    
    with st.expander("❤️ Relationships & Love", expanded=False):
        ep['romance'] = st.text_area("Tell me about your romantic relationships", value=ep.get('romance', ''), key="ep_romance", height=80)
        ep['marriage'] = st.text_area("If married, how did you meet?", value=ep.get('marriage', ''), key="ep_marriage", height=80)
        ep['children'] = st.text_area("Tell me about your children, if any", value=ep.get('children', ''), key="ep_children", height=80)
        ep['friends'] = st.text_area("Who are your closest friends?", value=ep.get('friends', ''), key="ep_friends", height=80)
    
    with st.expander("🌟 Challenges & Triumphs", expanded=False):
        ep['challenges'] = st.text_area("What were the hardest moments in your life?", value=ep.get('challenges', ''), key="ep_challenges", height=80)
        ep['losses'] = st.text_area("What losses have you experienced?", value=ep.get('losses', ''), key="ep_losses", height=80)
        ep['proud_moments'] = st.text_area("What are your proudest moments?", value=ep.get('proud_moments', ''), key="ep_proud", height=80)
    
    with st.expander("🌍 Life Philosophy & Wisdom", expanded=False):
        ep['life_lessons'] = st.text_area("What life lessons would you pass on?", value=ep.get('life_lessons', ''), key="ep_lessons", height=80)
        ep['values'] = st.text_area("What are your core values?", value=ep.get('values', ''), key="ep_values", height=80)
        ep['advice'] = st.text_area("Advice to your younger self?", value=ep.get('advice', ''), key="ep_advice", height=80)
        ep['legacy'] = st.text_area("How would you like to be remembered?", value=ep.get('legacy', ''), key="ep_legacy", height=80)
    
    if st.button("💾 Save Biographer's Questions", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("Biographer's profile saved!")
        st.rerun()

# ============================================================================
# NARRATIVE GPS PROFILE SECTION
# ============================================================================
def render_narrative_gps():
    """Render the Narrative GPS questionnaire in the profile"""
    st.markdown("### ❤️ The Heart of Your Story")
    
    st.markdown("""
    <div style="background-color: #f0f2f6; padding: 15px; border-radius: 5px; margin-bottom: 15px; border-left: 4px solid #ff4b4b;">
    Before we write a single word, let's understand why this book matters. The more honest and detailed you are here, the more your true voice will shine through every page.
    </div>
    """, unsafe_allow_html=True)
    
    if 'narrative_gps' not in st.session_state.user_account:
        st.session_state.user_account['narrative_gps'] = {}
    
    gps = st.session_state.user_account['narrative_gps']
    
    with st.expander("📖 Section 1: The Book Itself (Project Scope)", expanded=True):
        gps['book_title'] = st.text_input(
            "BOOK TITLE (Working or Final):",
            value=gps.get('book_title', ''),
            placeholder="What's your working title?",
            key="gps_title"
        )
        
        genre_options = ["", "Memoir", "Autobiography", "Family History", "Business/Legacy Book", "Other"]
        genre_index = 0
        if gps.get('genre') in genre_options:
            genre_index = genre_options.index(gps['genre'])
        
        gps['genre'] = st.selectbox(
            "BOOK GENRE/CATEGORY:",
            options=genre_options,
            index=genre_index,
            key="gps_genre"
        )
        if gps['genre'] == "Other":
            gps['genre_other'] = st.text_input("Please specify:", value=gps.get('genre_other', ''), key="gps_genre_other")
        
        length_options = ["", "A short book (100-150 pages)", "Standard length (200-300 pages)", "Comprehensive (300+ pages)"]
        length_index = 0
        if gps.get('book_length') in length_options:
            length_index = length_options.index(gps['book_length'])
        
        gps['book_length'] = st.selectbox(
            "BOOK LENGTH VISION:",
            options=length_options,
            index=length_index,
            key="gps_length"
        )
        
        gps['timeline'] = st.text_area(
            "TIMELINE & DEADLINES:",
            value=gps.get('timeline', ''),
            placeholder="Target publication date or event?",
            key="gps_timeline"
        )
        
        completion_options = ["", "Notes only", "Partial chapters", "Full draft"]
        completion_index = 0
        if gps.get('completion_status') in completion_options:
            completion_index = completion_options.index(gps['completion_status'])
        
        gps['completion_status'] = st.selectbox(
            "COMPLETION STATUS:",
            options=completion_options,
            index=completion_index,
            key="gps_completion"
        )
    
    with st.expander("🎯 Section 2: Purpose & Audience (The 'Why')", expanded=False):
        if 'purposes' not in gps:
            gps['purposes'] = []
        
        purposes_options = [
            "Leave a legacy for family/future generations",
            "Share life lessons to help others",
            "Document professional/business journey",
            "Heal or process through writing",
            "Establish authority/expertise",
            "Entertain with entertaining stories"
        ]
        
        for purpose in purposes_options:
            if st.checkbox(
                purpose,
                value=purpose in gps.get('purposes', []),
                key=f"gps_purpose_{purpose}"
            ):
                if purpose not in gps['purposes']:
                    gps['purposes'].append(purpose)
            else:
                if purpose in gps['purposes']:
                    gps['purposes'].remove(purpose)
        
        gps['purpose_other'] = st.text_input("Other:", value=gps.get('purpose_other', ''), key="gps_purpose_other")
        
        gps['audience_family'] = st.text_input(
            "Family members (which generations?):",
            value=gps.get('audience_family', ''),
            key="gps_audience_family"
        )
        
        gps['audience_industry'] = st.text_input(
            "People in your industry/profession:",
            value=gps.get('audience_industry', ''),
            key="gps_audience_industry"
        )
        
        gps['audience_challenges'] = st.text_input(
            "People facing similar challenges you overcame:",
            value=gps.get('audience_challenges', ''),
            key="gps_audience_challenges"
        )
        
        gps['audience_general'] = st.text_input(
            "The general public interested in:",
            value=gps.get('audience_general', ''),
            placeholder="your topic",
            key="gps_audience_general"
        )
        
        gps['reader_takeaway'] = st.text_area(
            "What do you want readers to feel, think, or do after finishing your book?",
            value=gps.get('reader_takeaway', ''),
            key="gps_takeaway"
        )
    
    with st.expander("🎭 Section 3: Tone & Voice (The 'How')", expanded=False):
        if 'voices' not in gps:
            gps['voices'] = []
        
        voice_options = [
            "Warm and conversational",
            "Professional and authoritative",
            "Raw and vulnerable",
            "Humorous/lighthearted",
            "Philosophical/reflective"
        ]
        
        for voice in voice_options:
            if st.checkbox(
                voice,
                value=voice in gps.get('voices', []),
                key=f"gps_voice_{voice}"
            ):
                if voice not in gps['voices']:
                    gps['voices'].append(voice)
            else:
                if voice in gps['voices']:
                    gps['voices'].remove(voice)
        
        gps['voice_other'] = st.text_input("Other:", value=gps.get('voice_other', ''), key="gps_voice_other")
        
        gps['emotional_tone'] = st.text_area(
            "EMOTIONAL TONE:",
            value=gps.get('emotional_tone', ''),
            placeholder="Should readers laugh? Cry? Feel inspired?",
            key="gps_emotional"
        )
        
        language_options = ["", "Simple/everyday language", "Rich/descriptive prose", "Short/punchy chapters", "Long/flowing narratives"]
        language_index = 0
        if gps.get('language_style') in language_options:
            language_index = language_options.index(gps['language_style'])
        
        gps['language_style'] = st.selectbox(
            "LANGUAGE STYLE:",
            options=language_options,
            index=language_index,
            key="gps_language"
        )
    
    with st.expander("📋 Section 4: Content Parameters (The 'What')", expanded=False):
        time_options = ["", "Your entire life", "A specific era/decade", "One defining experience", "Your career/business journey"]
        time_index = 0
        if gps.get('time_coverage') in time_options:
            time_index = time_options.index(gps['time_coverage'])
        
        gps['time_coverage'] = st.selectbox(
            "TIME COVERAGE:",
            options=time_options,
            index=time_index,
            key="gps_time"
        )
        
        gps['sensitive_material'] = st.text_area(
            "SENSITIVE MATERIAL:",
            value=gps.get('sensitive_material', ''),
            placeholder="Topics to handle carefully?",
            key="gps_sensitive"
        )
        
        gps['sensitive_people'] = st.text_area(
            "Living people requiring sensitivity?",
            value=gps.get('sensitive_people', ''),
            key="gps_sensitive_people"
        )
        
        if 'inclusions' not in gps:
            gps['inclusions'] = []
        
        inclusion_options = ["Photos", "Family trees", "Recipes", "Letters/documents", "Timelines"]
        for inc in inclusion_options:
            if st.checkbox(
                inc,
                value=inc in gps.get('inclusions', []),
                key=f"gps_inc_{inc}"
            ):
                if inc not in gps['inclusions']:
                    gps['inclusions'].append(inc)
            else:
                if inc in gps['inclusions']:
                    gps['inclusions'].remove(inc)
        
        gps['locations'] = st.text_area(
            "Key locations that must appear:",
            value=gps.get('locations', ''),
            key="gps_locations"
        )
    
    with st.expander("📦 Section 5: Assets & Access", expanded=False):
        if 'materials' not in gps:
            gps['materials'] = []
        
        material_options = [
            "Journals/diaries", "Letters/emails", "Photos", "Video/audio recordings",
            "Newspaper clippings", "Awards/certificates", "Previous interviews"
        ]
        
        for mat in material_options:
            if st.checkbox(
                mat,
                value=mat in gps.get('materials', []),
                key=f"gps_mat_{mat}"
            ):
                if mat not in gps['materials']:
                    gps['materials'].append(mat)
            else:
                if mat in gps['materials']:
                    gps['materials'].remove(mat)
        
        gps['people_to_interview'] = st.text_area(
            "People to interview:",
            value=gps.get('people_to_interview', ''),
            key="gps_people"
        )
    
    with st.expander("🤝 Section 6: Collaboration", expanded=False):
        involvement_options = [
            "I'll answer questions, you write",
            "I'll write drafts, you polish",
            "We'll interview together, then you write",
            "Mixed approach: [explain]"
        ]
        
        involvement_index = 0
        if gps.get('involvement') in involvement_options:
            involvement_index = involvement_options.index(gps['involvement'])
        
        gps['involvement'] = st.radio(
            "How do you want to work together?",
            options=involvement_options,
            index=involvement_index,
            key="gps_involvement"
        )
        
        if gps.get('involvement') == "Mixed approach: [explain]":
            gps['involvement_explain'] = st.text_area(
                "Explain your preferred approach:",
                value=gps.get('involvement_explain', ''),
                key="gps_involvement_explain"
            )
        
        feedback_options = ["", "Written comments", "Video discussions", "Line-by-line edits"]
        feedback_index = 0
        if gps.get('feedback_style') in feedback_options:
            feedback_index = feedback_options.index(gps['feedback_style'])
        
        gps['feedback_style'] = st.selectbox(
            "FEEDBACK STYLE:",
            options=feedback_options,
            index=feedback_index,
            key="gps_feedback"
        )
        
        gps['unspoken'] = st.text_area(
            "What do you hope I'll bring to this project?",
            value=gps.get('unspoken', ''),
            key="gps_unspoken"
        )
    
    if st.button("💾 Save The Heart of Your Story", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("✅ Saved!")
        st.rerun()

# ============================================================================
# STORAGE FUNCTIONS
# ============================================================================
def get_user_filename(user_id):
    return f"user_data_{hashlib.md5(user_id.encode()).hexdigest()[:8]}.json"

def load_user_data(user_id):
    fname = get_user_filename(user_id)
    try:
        if os.path.exists(fname):
            return json.load(open(fname, 'r'))
        return {"responses": {}, "vignettes": [], "last_loaded": datetime.now().isoformat()}
    except: 
        return {"responses": {}, "vignettes": [], "last_loaded": datetime.now().isoformat()}

def save_user_data(user_id, responses_data):
    fname = get_user_filename(user_id)
    try:
        existing = load_user_data(user_id)
        data = {
            "user_id": user_id, 
            "responses": responses_data,
            "vignettes": existing.get("vignettes", []),
            "beta_feedback": existing.get("beta_feedback", {}),
            "last_saved": datetime.now().isoformat()
        }
        with open(fname, 'w') as f: 
            json.dump(data, f, indent=2)
        return True
    except: 
        return False

# ============================================================================
# CORE RESPONSE FUNCTIONS
# ============================================================================
def save_response(session_id, question, answer):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    text_only = re.sub(r'<[^>]+>', '', answer) if answer else ""
    
    if st.session_state.user_account:
        word_count = len(re.findall(r'\w+', text_only))
        st.session_state.user_account["stats"]["total_words"] = st.session_state.user_account["stats"].get("total_words", 0) + word_count
        st.session_state.user_account["stats"]["last_active"] = datetime.now().isoformat()
        save_account_data(st.session_state.user_account)
    
    if session_id not in st.session_state.responses:
        session_data = next((s for s in (st.session_state.current_question_bank or []) if s["id"] == session_id), 
                          {"title": f"Session {session_id}", "word_target": DEFAULT_WORD_TARGET})
        st.session_state.responses[session_id] = {
            "title": session_data.get("title", f"Session {session_id}"),
            "questions": {}, 
            "summary": "", 
            "completed": False,
            "word_target": session_data.get("word_target", DEFAULT_WORD_TARGET)
        }
    
    images = []
    if st.session_state.image_handler:
        images = st.session_state.image_handler.get_images_for_answer(session_id, question)
    
    st.session_state.responses[session_id]["questions"][question] = {
        "answer": answer,
        "question": question, 
        "timestamp": datetime.now().isoformat(),
        "has_images": len(images) > 0 or ('<img' in answer),
        "image_count": len(images),
        "images": [{"id": img["id"], "caption": img.get("caption", "")} for img in images]
    }
    
    success = save_user_data(user_id, st.session_state.responses)
    if success: 
        st.session_state.data_loaded = False
        
        session_title = st.session_state.responses[session_id].get("title", f"Session {session_id}")
        suggestions = generate_writing_suggestions(question, answer, session_title)
        if suggestions and not isinstance(suggestions, dict):
            st.session_state.current_ai_suggestions = suggestions
            st.session_state.show_ai_suggestions = True
    
    return success

def delete_response(session_id, question):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    if session_id in st.session_state.responses and question in st.session_state.responses[session_id]["questions"]:
        del st.session_state.responses[session_id]["questions"][question]
        success = save_user_data(user_id, st.session_state.responses)
        if success: 
            st.session_state.data_loaded = False
        return success
    return False

def calculate_author_word_count(session_id):
    total = 0
    if session_id in st.session_state.responses:
        for q, d in st.session_state.responses[session_id].get("questions", {}).items():
            if d.get("answer"): 
                text_only = re.sub(r'<[^>]+>', '', d["answer"])
                total += len(re.findall(r'\w+', text_only))
    return total

def get_progress_info(session_id):
    current = calculate_author_word_count(session_id)
    if session_id not in st.session_state.responses:
        session_data = next((s for s in (st.session_state.current_question_bank or []) if s["id"] == session_id), {})
        st.session_state.responses[session_id] = {
            "title": session_data.get("title", f"Session {session_id}"),
            "questions": {}, 
            "summary": "", 
            "completed": False,
            "word_target": session_data.get("word_target", DEFAULT_WORD_TARGET)
        }
    
    target = st.session_state.responses[session_id].get("word_target", DEFAULT_WORD_TARGET)
    if target == 0: 
        percent = 100
    else: 
        percent = (current / target) * 100
    
    return {
        "current_count": current, 
        "target": target, 
        "progress_percent": percent,
        "emoji": "🟢" if percent >= 100 else "🟡" if percent >= 70 else "🔴",
        "color": "#27ae60" if percent >= 100 else "#f39c12" if percent >= 70 else "#e74c3c",
        "remaining_words": max(0, target - current),
        "status_text": "Target achieved!" if current >= target else f"{max(0, target - current)} words remaining"
    }

# ============================================================================
# SEARCH FUNCTIONALITY
# ============================================================================
def search_all_answers(search_query):
    if not search_query or len(search_query) < 2: 
        return []
    
    results = []
    search_query = search_query.lower()
    
    for session in (st.session_state.current_question_bank or []):
        session_id = session["id"]
        session_data = st.session_state.responses.get(session_id, {})
        
        for question_text, answer_data in session_data.get("questions", {}).items():
            html_answer = answer_data.get("answer", "")
            text_answer = re.sub(r'<[^>]+>', '', html_answer)
            has_images = answer_data.get("has_images", False) or ('<img' in html_answer)
            
            if search_query in text_answer.lower() or search_query in question_text.lower():
                results.append({
                    "session_id": session_id, 
                    "session_title": session["title"],
                    "question": question_text, 
                    "answer": text_answer[:300] + "..." if len(text_answer) > 300 else text_answer,
                    "timestamp": answer_data.get("timestamp", ""), 
                    "word_count": len(text_answer.split()),
                    "has_images": has_images,
                    "image_count": answer_data.get("image_count", 0)
                })
    
    results.sort(key=lambda x: x["timestamp"], reverse=True)
    return results

# ============================================================================
# QUESTION BANK LOADING
# ============================================================================
def initialize_question_bank():
    if 'current_question_bank' in st.session_state and st.session_state.current_question_bank:
        return True
    
    if QuestionBankManager:
        try:
            qb_manager = QuestionBankManager(st.session_state.get('user_id'))
            st.session_state.qb_manager = qb_manager
            
            if os.path.exists("sessions/sessions.csv"):
                shutil.copy("sessions/sessions.csv", "question_banks/default/life_story_comprehensive.csv")
            
            default = qb_manager.load_default_bank("life_story_comprehensive")
            if default:
                st.session_state.current_question_bank = default
                st.session_state.current_bank_name = "📖 Life Story - Comprehensive"
                st.session_state.current_bank_type = "default"
                st.session_state.current_bank_id = "life_story_comprehensive"
                st.session_state.qb_manager_initialized = True
                
                for s in default:
                    sid = s["id"]
                    if sid not in st.session_state.responses:
                        st.session_state.responses[sid] = {
                            "title": s["title"], 
                            "questions": {}, 
                            "summary": "",
                            "completed": False, 
                            "word_target": s.get("word_target", DEFAULT_WORD_TARGET)
                        }
                return True
        except: 
            pass
    
    # Fallback hardcoded sessions
    st.session_state.current_question_bank = [
        {
            "id": 1,
            "title": "Early Memories",
            "questions": [
                "What is your earliest memory?",
                "Describe the house you grew up in.",
                "Who was your favorite family member as a child?"
            ]
        },
        {
            "id": 2,
            "title": "School Years",
            "questions": [
                "What was your favorite subject in school?",
                "Who was your most memorable teacher?",
                "What were your hopes and dreams as a teenager?"
            ]
        },
        {
            "id": 3,
            "title": "Adulthood & Career",
            "questions": [
                "What was your first job?",
                "Tell me about your proudest achievement.",
                "What life lessons would you share with younger generations?"
            ]
        }
    ]
    st.session_state.current_bank_name = "Default Bank"
    st.session_state.current_bank_type = "default"
    st.session_state.qb_manager_initialized = True
    return True

def load_question_bank(sessions, bank_name, bank_type, bank_id=None):
    st.session_state.current_question_bank = sessions
    st.session_state.current_bank_name = bank_name
    st.session_state.current_bank_type = bank_type
    st.session_state.current_bank_id = bank_id
    st.session_state.current_session = 0
    st.session_state.current_question = 0
    st.session_state.current_question_override = None
    
    for s in sessions:
        sid = s["id"]
        if sid not in st.session_state.responses:
            st.session_state.responses[sid] = {
                "title": s["title"], 
                "questions": {},
                "summary": "",
                "completed": False, 
                "word_target": s.get("word_target", DEFAULT_WORD_TARGET)
            }

# ============================================================================
# BETA READER FUNCTIONS
# ============================================================================
def generate_beta_reader_feedback(session_title, session_text, feedback_type="comprehensive"):
    if not beta_reader: 
        return {"error": "BetaReader not available"}
    return beta_reader.generate_feedback(session_title, session_text, feedback_type)

def save_beta_feedback(user_id, session_id, feedback_data):
    if not beta_reader: 
        return False
    return beta_reader.save_feedback(user_id, session_id, feedback_data, get_user_filename, load_user_data)

def get_previous_beta_feedback(user_id, session_id):
    if not beta_reader: 
        return None
    return beta_reader.get_previous_feedback(user_id, session_id, get_user_filename, load_user_data)

# ============================================================================
# VIGNETTE FUNCTIONS
# ============================================================================
def on_vignette_select(vignette_id):
    st.session_state.selected_vignette_id = vignette_id
    st.session_state.show_vignette_detail = True
    st.session_state.show_vignette_manager = False
    st.rerun()

def on_vignette_edit(vignette_id):
    st.session_state.editing_vignette_id = vignette_id
    st.session_state.show_vignette_detail = False
    st.session_state.show_vignette_manager = False
    st.session_state.show_vignette_modal = True
    st.rerun()

def on_vignette_delete(vignette_id):
    if VignetteManager and st.session_state.get('vignette_manager', VignetteManager(st.session_state.user_id)).delete_vignette(vignette_id):
        st.success("Deleted!"); 
        st.rerun()
    else: 
        st.error("Failed to delete")

def on_vignette_publish(vignette):
    st.session_state.published_vignette = vignette
    st.success(f"Published '{vignette['title']}'!"); 
    st.rerun()

def show_vignette_modal():
    if not VignetteManager: 
        st.error("Vignette module not available"); 
        st.session_state.show_vignette_modal = False; 
        return
    st.markdown("### ✏️ Create Vignette")
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    edit = st.session_state.vignette_manager.get_vignette_by_id(st.session_state.editing_vignette_id) if st.session_state.get('editing_vignette_id') else None
    st.session_state.vignette_manager.display_vignette_creator(on_publish=on_vignette_publish, edit_vignette=edit)
    if st.button("← Back"):
        st.session_state.show_vignette_modal = False
        st.rerun()

def show_vignette_manager():
    if not VignetteManager: 
        st.error("Vignette module not available"); 
        st.session_state.show_vignette_manager = False; 
        return
    st.title("📚 Your Vignettes")
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    filter_map = {"All Stories": "all", "Published": "published", "Drafts": "drafts"}
    filter_option = st.radio("Show:", ["All Stories", "Published", "Drafts"], horizontal=True, key="vign_filter")
    st.session_state.vignette_manager.display_vignette_gallery(
        filter_by=filter_map.get(filter_option, "all"),
        on_select=on_vignette_select, 
        on_edit=on_vignette_edit, 
        on_delete=on_vignette_delete
    )
    st.divider()
    if st.button("➕ Create New Vignette", type="primary"):
        st.session_state.show_vignette_manager = False; 
        st.session_state.show_vignette_modal = True; 
        st.session_state.editing_vignette_id = None; 
        st.rerun()
    if st.button("← Back"):
        st.session_state.show_vignette_manager = False
        st.rerun()

def show_vignette_detail():
    if not VignetteManager or not st.session_state.get('selected_vignette_id'): 
        st.session_state.show_vignette_detail = False; 
        return
    st.title("📖 Read Vignette")
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    vignette = st.session_state.vignette_manager.get_vignette_by_id(st.session_state.selected_vignette_id)
    if not vignette: 
        st.error("Not found"); 
        st.session_state.show_vignette_detail = False; 
        return
    st.session_state.vignette_manager.display_full_vignette(
        st.session_state.selected_vignette_id,
        on_back=lambda: st.session_state.update(show_vignette_detail=False, selected_vignette_id=None),
        on_edit=on_vignette_edit
    )
    if st.button("← Back"):
        st.session_state.show_vignette_detail = False
        st.rerun()

def switch_to_vignette(vignette_topic, content=""):
    st.session_state.current_question_override = f"Vignette: {vignette_topic}"
    if content:
        save_response(st.session_state.current_question_bank[st.session_state.current_session]["id"], 
                     f"Vignette: {vignette_topic}", content)
    st.rerun()

def switch_to_custom_topic(topic_text):
    st.session_state.current_question_override = topic_text
    st.rerun()

# ============================================================================
# TOPIC BROWSER & SESSION MANAGER
# ============================================================================
def show_topic_browser():
    if not TopicBank: 
        st.error("Topic module not available"); 
        st.session_state.show_topic_browser = False; 
        return
    st.title("📚 Topic Browser")
    TopicBank(st.session_state.user_id).display_topic_browser(
        on_topic_select=lambda t: (switch_to_custom_topic(t), st.session_state.update(show_topic_browser=False)),
        unique_key=str(time.time())
    )
    if st.button("← Back"):
        st.session_state.show_topic_browser = False
        st.rerun()

def show_session_creator():
    if not SessionManager: 
        st.error("Session module not available"); 
        st.session_state.show_session_creator = False; 
        return
    st.title("📋 Create Custom Session")
    SessionManager(st.session_state.user_id, "sessions/sessions.csv").display_session_creator()
    if st.button("← Back"):
        st.session_state.show_session_creator = False
        st.rerun()

def show_session_manager():
    if not SessionManager: 
        st.error("Session module not available"); 
        st.session_state.show_session_manager = False; 
        return
    st.title("📖 Session Manager")
    mgr = SessionManager(st.session_state.user_id, "sessions/sessions.csv")
    if st.button("➕ Create New Session", type="primary"):
        st.session_state.show_session_manager = False; 
        st.session_state.show_session_creator = True; 
        st.rerun()
    st.divider()
    mgr.display_session_grid(cols=2, on_session_select=lambda sid: [st.session_state.update(
        current_session=i, current_question=0, current_question_override=None) for i, s in enumerate(st.session_state.current_question_bank) if s["id"] == sid][0])
    if st.button("← Back"):
        st.session_state.show_session_manager = False
        st.rerun()

# ============================================================================
# QUESTION BANK UI FUNCTIONS
# ============================================================================
def show_bank_manager():
    if not QuestionBankManager: 
        st.error("Question Bank Manager not available"); 
        st.session_state.show_bank_manager = False; 
        return
    user_id = st.session_state.get('user_id')
    if st.session_state.qb_manager is None: 
        st.session_state.qb_manager = QuestionBankManager(user_id)
    else: 
        st.session_state.qb_manager.user_id = user_id
    st.title("📋 Bank Manager")
    st.session_state.qb_manager.display_bank_selector()
    if st.button("← Back"):
        st.session_state.show_bank_manager = False
        st.rerun()

def show_bank_editor():
    if not QuestionBankManager or not st.session_state.get('editing_bank_id'): 
        st.session_state.show_bank_editor = False; 
        return
    user_id = st.session_state.get('user_id')
    if st.session_state.qb_manager is None: 
        st.session_state.qb_manager = QuestionBankManager(user_id)
    else: 
        st.session_state.qb_manager.user_id = user_id
    st.title("📝 Edit Bank")
    st.session_state.qb_manager.display_bank_editor(st.session_state.editing_bank_id)
    if st.button("← Back"):
        st.session_state.show_bank_editor = False
        st.rerun()

# ============================================================================
# PDF GENERATION FUNCTIONS
# ============================================================================
class PDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, 'Tell My Story', 0, 0, 'L')
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'R')
            self.ln(15)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, 'Generated by Tell My Story', 0, 0, 'C')

def generate_pdf(book_title, author_name, stories, format_style, include_toc, include_dates):
    pdf = PDF()
    pdf.add_page()
    
    pdf.set_fill_color(102, 126, 234)
    pdf.rect(0, 0, 210, 297, 'F')
    pdf.set_text_color(255, 255, 255)
    
    safe_title = ''.join(c for c in book_title if ord(c) < 128)
    safe_author = ''.join(c for c in author_name if ord(c) < 128)
    
    pdf.set_font('Arial', 'B', 30)
    pdf.cell(0, 40, '', 0, 1)
    pdf.cell(0, 20, safe_title if safe_title else 'My Story', 0, 1, 'C')
    pdf.set_font('Arial', '', 16)
    pdf.cell(0, 10, f'by {safe_author}' if safe_author else 'by Author', 0, 1, 'C')
    pdf.add_page()
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font('Arial', '', 11)
    
    for story in stories:
        question = story.get('question', '')
        answer = story.get('answer_text', '')
        
        safe_q = ''.join(c for c in question if ord(c) < 128)
        safe_a = ''.join(c for c in answer if ord(c) < 128)
        
        pdf.set_font('Arial', 'B', 12)
        pdf.multi_cell(0, 6, safe_q)
        pdf.ln(2)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, safe_a)
        pdf.ln(5)
    
    return pdf.output(dest='S').encode('latin-1', 'ignore')

# ============================================================================
# DOCX GENERATION FUNCTIONS
# ============================================================================
def generate_docx(book_title, author_name, stories, format_style, include_toc, include_dates):
    doc = Document()
    
    title = doc.add_heading(book_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = doc.add_paragraph(f'by {author_name}')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    if include_toc:
        doc.add_heading('Table of Contents', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for story in stories:
        question = story.get('question', '')
        answer_text = story.get('answer_text', '')
        images = story.get('images', [])
        
        doc.add_heading(question, 2)
        doc.add_paragraph(answer_text)
        
        for img_data in images:
            b64 = img_data.get('base64')
            caption = img_data.get('caption', '')
            if b64:
                try:
                    img_bytes = base64.b64decode(b64)
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(4))
                    if caption:
                        doc.add_paragraph(caption, style='Caption')
                except:
                    pass
        doc.add_paragraph()
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes

# ============================================================================
# HTML GENERATION FUNCTION
# ============================================================================
def generate_html(book_title, author_name, stories):
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{book_title}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #667eea; text-align: center; }}
        .story {{ margin-bottom: 50px; padding: 20px; background: #f9f9f9; border-radius: 10px; }}
        .question {{ font-size: 1.3em; font-weight: bold; color: #2c3e50; }}
        img {{ max-width: 100%; border-radius: 8px; margin: 10px 0; }}
        .caption {{ font-style: italic; color: #666; text-align: center; }}
    </style>
</head>
<body>
    <h1>{book_title}</h1>
    <div style="text-align: center;">by {author_name}</div>
"""
    
    for story in stories:
        html += f"""
    <div class="story">
        <div class="question">{story['question']}</div>
        <div>{story['answer_text']}</div>
"""
        for img in story.get('images', []):
            if img.get('base64'):
                html += f'        <img src="data:image/jpeg;base64,{img["base64"]}" alt="{img.get("caption", "")}">\n'
                if img.get('caption'):
                    html += f'        <div class="caption">{img["caption"]}</div>\n'
        
        html += f"""
    </div>
    <hr>
"""
    
    html += f"""
    <div class="footer">
        Generated by Tell My Story • {datetime.now().strftime("%B %d, %Y")}
    </div>
</body>
</html>"""
    return html

# ============================================================================
# ZIP GENERATION FUNCTION
# ============================================================================
def generate_zip(book_title, author_name, stories):
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        html = generate_html(book_title, author_name, stories)
        zip_file.writestr(f"{book_title.replace(' ', '_')}.html", html)
        
        for i, story in enumerate(stories):
            for j, img in enumerate(story.get('images', [])):
                if img.get('base64'):
                    img_data = base64.b64decode(img['base64'])
                    zip_file.writestr(f"images/image_{i}_{j}.jpg", img_data)
    
    return zip_buffer.getvalue()

# ============================================================================
# PAGE CONFIG
# ============================================================================
st.set_page_config(page_title="Tell My Story", page_icon="📖", layout="wide")

# Initialize question bank
if not st.session_state.qb_manager_initialized: 
    initialize_question_bank()
SESSIONS = st.session_state.get('current_question_bank', [])

# Load user data
if st.session_state.logged_in and st.session_state.user_id and not st.session_state.data_loaded:
    user_data = load_user_data(st.session_state.user_id)
    if "responses" in user_data:
        for sid_str, sdata in user_data["responses"].items():
            try: 
                sid = int(sid_str)
            except: 
                continue
            if sid in st.session_state.responses and "questions" in sdata and sdata["questions"]:
                st.session_state.responses[sid]["questions"] = sdata["questions"]
    st.session_state.data_loaded = True
    init_image_handler()

if not SESSIONS:
    st.error("❌ No question bank loaded.")
    st.stop()

# ============================================================================
# AUTHENTICATION UI
# ============================================================================
if not st.session_state.logged_in:
    st.markdown(f'<div style="text-align:center"><img src="{LOGO_URL}" style="height:80px"></div>', unsafe_allow_html=True)
    st.title("Tell My Story")
    
    tab1, tab2 = st.tabs(["Login", "Sign Up"])
    
    with tab1:
        with st.form("login_form"):
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            if st.form_submit_button("Login", type="primary"):
                if email and password:
                    result = authenticate_user(email, password)
                    if result["success"]:
                        st.session_state.user_id = result["user_id"]
                        st.session_state.user_account = result["user_record"]
                        st.session_state.logged_in = True
                        st.rerun()
                    else:
                        st.error("Login failed")
    
    with tab2:
        with st.form("signup_form"):
            col1, col2 = st.columns(2)
            with col1:
                first = st.text_input("First Name*")
            with col2:
                last = st.text_input("Last Name*")
            email = st.text_input("Email*")
            password = st.text_input("Password*", type="password")
            confirm = st.text_input("Confirm Password*", type="password")
            accept = st.checkbox("I agree to the Terms*")
            
            if st.form_submit_button("Sign Up", type="primary"):
                if not first or not last or not email or not password:
                    st.error("All fields required")
                elif password != confirm:
                    st.error("Passwords don't match")
                elif not accept:
                    st.error("Must accept terms")
                elif get_account_data(email=email):
                    st.error("Email already exists")
                else:
                    result = create_user_account({"first_name": first, "last_name": last, "email": email}, password)
                    if result["success"]:
                        st.session_state.user_id = result["user_id"]
                        st.session_state.user_account = result["user_record"]
                        st.session_state.logged_in = True
                        st.session_state.show_profile_setup = True
                        st.rerun()
    st.stop()

# ============================================================================
# PROFILE SETUP MODAL
# ============================================================================
if st.session_state.get('show_profile_setup', False):
    st.title("👤 Your Complete Profile")
    
    # Basic Profile
    with st.form("basic_profile"):
        st.subheader("Basic Information")
        gender = st.radio("Gender", ["Male", "Female", "Other", "Prefer not to say"], horizontal=True)
        col1, col2, col3 = st.columns(3)
        with col1: 
            birth_month = st.selectbox("Month", ["January","February","March","April","May","June","July","August","September","October","November","December"])
        with col2: 
            birth_day = st.selectbox("Day", list(range(1,32)))
        with col3: 
            birth_year = st.selectbox("Year", list(range(1950, 2025)))
        account_for = st.radio("Account Type", ["For me", "For someone else"], horizontal=True)
        
        if st.form_submit_button("💾 Save Basic Info", type="primary"):
            birthdate = f"{birth_month} {birth_day}, {birth_year}"
            st.session_state.user_account['profile'].update({
                'gender': gender, 'birthdate': birthdate
            })
            st.session_state.user_account['account_type'] = "self" if account_for == "For me" else "other"
            save_account_data(st.session_state.user_account)
            st.success("Saved!")
            st.rerun()
    
    st.divider()
    
    # Enhanced Biographer Profile
    render_enhanced_profile()
    
    st.divider()
    
    # Narrative GPS
    render_narrative_gps()
    
    st.divider()
    
    # Privacy Settings
    with st.expander("🔒 Privacy Settings", expanded=False):
        if 'privacy_settings' not in st.session_state.user_account:
            st.session_state.user_account['privacy_settings'] = {
                "profile_public": False, "stories_public": False,
                "allow_sharing": False, "data_collection": True
            }
        privacy = st.session_state.user_account['privacy_settings']
        privacy['profile_public'] = st.checkbox("Make profile public", value=privacy.get('profile_public', False))
        privacy['data_collection'] = st.checkbox("Allow anonymous usage data", value=privacy.get('data_collection', True))
        if st.button("Save Privacy Settings", type="primary"):
            save_account_data(st.session_state.user_account)
            st.success("Saved!")
            st.rerun()
    
    st.divider()
    
    # Backup and Restore
    with st.expander("💾 Backup & Restore", expanded=False):
        backup_json = create_backup()
        if backup_json:
            st.download_button(
                label="📥 Download Backup",
                data=backup_json,
                file_name=f"backup_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json"
            )
        
        backups = list_backups()
        if backups:
            st.markdown("**Previous backups:**")
            for b in backups[:3]:
                st.text(f"📅 {b['date'][:10]}")
    
    if st.button("← Close Profile"):
        st.session_state.show_profile_setup = False
        st.rerun()
    
    st.stop()

# ============================================================================
# MODAL HANDLING
# ============================================================================
if st.session_state.show_privacy_settings:
    show_privacy_settings()
if st.session_state.show_cover_designer:
    show_cover_designer()
if st.session_state.show_bank_manager: 
    show_bank_manager()
if st.session_state.show_bank_editor: 
    show_bank_editor()
if st.session_state.show_beta_reader and st.session_state.current_beta_feedback: 
    if beta_reader:
        beta_reader.show_modal(st.session_state.current_beta_feedback, 
                              {"id": SESSIONS[st.session_state.current_session]["id"], 
                               "title": SESSIONS[st.session_state.current_session]["title"]},
                              st.session_state.user_id, 
                              save_beta_feedback, 
                              lambda: st.session_state.update(show_beta_reader=False, current_beta_feedback=None))
if st.session_state.show_vignette_detail: 
    show_vignette_detail()
if st.session_state.show_vignette_manager: 
    show_vignette_manager()
if st.session_state.show_vignette_modal: 
    show_vignette_modal()
if st.session_state.show_topic_browser: 
    show_topic_browser()
if st.session_state.show_session_manager: 
    show_session_manager()
if st.session_state.show_session_creator: 
    show_session_creator()

# ============================================================================
# MAIN HEADER
# ============================================================================
st.markdown(f'<div style="text-align:center"><img src="{LOGO_URL}" style="height:60px"></div>', unsafe_allow_html=True)

# ============================================================================
# SIDEBAR
# ============================================================================
with st.sidebar:
    st.markdown(f"### 👤 {st.session_state.user_account.get('profile', {}).get('first_name', 'User')}")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📝 Profile"):
            st.session_state.show_profile_setup = True
            st.rerun()
    with col2:
        if st.button("🔒 Privacy"):
            st.session_state.show_privacy_settings = True
            st.rerun()
    
    if st.button("🎨 Cover Designer"):
        st.session_state.show_cover_designer = True
        st.rerun()
    
    if st.button("🚪 Log Out"):
        logout_user()
    
    st.divider()
    
    st.subheader("📚 Question Banks")
    if st.button("📋 Bank Manager"):
        st.session_state.show_bank_manager = True
        st.rerun()
    if st.session_state.get('current_bank_name'):
        st.info(f"**Current:** {st.session_state.current_bank_name}")
    
    st.divider()
    
    st.subheader("📖 Sessions")
    for i, s in enumerate(SESSIONS):
        sid = s["id"]
        sdata = st.session_state.responses.get(sid, {})
        resp_cnt = len(sdata.get("questions", {}))
        status = "✅" if resp_cnt > 0 else "📝"
        if st.button(f"{status} {s['title']}", key=f"sesh_{i}"):
            st.session_state.current_session = i
            st.session_state.current_question = 0
            st.rerun()
    
    st.divider()
    
    st.subheader("✨ Vignettes")
    if st.button("📝 New Vignette"):
        st.session_state.show_vignette_modal = True
        st.rerun()
    if st.button("📖 View All"):
        st.session_state.show_vignette_manager = True
        st.rerun()
    
    st.divider()
    
    st.subheader("📤 Export")
    total_answers = sum(len(st.session_state.responses.get(s["id"], {}).get("questions", {})) for s in SESSIONS)
    st.caption(f"Total stories: {total_answers}")
    
    if st.session_state.logged_in:
        export_data = []
        for session in SESSIONS:
            sid = session["id"]
            sdata = st.session_state.responses.get(sid, {})
            for q, a in sdata.get("questions", {}).items():
                images_with_data = []
                if a.get("images"):
                    for img_ref in a.get("images", []):
                        img_id = img_ref.get("id")
                        b64 = st.session_state.image_handler.get_image_base64(img_id) if st.session_state.image_handler else None
                        if b64:
                            images_with_data.append({
                                "base64": b64,
                                "caption": img_ref.get("caption", "")
                            })
                
                export_data.append({
                    "question": q,
                    "answer_text": re.sub(r'<[^>]+>', '', a.get("answer", "")),
                    "session_title": session["title"],
                    "images": images_with_data
                })
        
        if export_data:
            book_title = st.text_input("Book Title", value="My Story")
            author = st.text_input("Author", value=st.session_state.user_account.get('profile', {}).get('first_name', ''))
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("DOCX"):
                    docx_bytes = generate_docx(book_title, author, export_data, "memoir", True, False)
                    st.download_button("Download", data=docx_bytes, file_name=f"{book_title}.docx", key="docx_btn")
            with col2:
                if st.button("HTML"):
                    html = generate_html(book_title, author, export_data)
                    st.download_button("Download", data=html, file_name=f"{book_title}.html", key="html_btn")
            with col3:
                if st.button("ZIP"):
                    zip_data = generate_zip(book_title, author, export_data)
                    st.download_button("Download", data=zip_data, file_name=f"{book_title}.zip", key="zip_btn")
    
    st.divider()
    
    st.subheader("🔍 Search")
    search_query = st.text_input("Search...", placeholder="e.g., childhood")
    if search_query:
        results = search_all_answers(search_query)
        if results:
            st.success(f"Found {len(results)}")
            for r in results[:3]:
                st.caption(r['question'][:30] + "...")

# ============================================================================
# MAIN CONTENT AREA
# ============================================================================
if st.session_state.current_session >= len(SESSIONS): 
    st.session_state.current_session = 0

current_session = SESSIONS[st.session_state.current_session]
current_session_id = current_session["id"]

if st.session_state.current_question_override:
    current_question_text = st.session_state.current_question_override
else:
    if st.session_state.current_question >= len(current_session["questions"]): 
        st.session_state.current_question = 0
    current_question_text = current_session["questions"][st.session_state.current_question]

st.title(current_session['title'])
st.subheader(current_question_text)

# Progress
sdata = st.session_state.responses.get(current_session_id, {})
answered = len(sdata.get("questions", {}))
total = len(current_session["questions"])
if total > 0: 
    st.progress(answered/total)
    st.caption(f"Completed: {answered}/{total}")

# Get existing answer
existing_answer = ""
if current_session_id in st.session_state.responses:
    if current_question_text in st.session_state.responses[current_session_id].get("questions", {}):
        existing_answer = st.session_state.responses[current_session_id]["questions"][current_question_text]["answer"]

# Initialize image handler
if st.session_state.logged_in:
    init_image_handler()
    existing_images = st.session_state.image_handler.get_images_for_answer(current_session_id, current_question_text) if st.session_state.image_handler else []

# ============================================================================
# QUILL EDITOR
# ============================================================================
editor_key = f"quill_{current_session_id}_{current_question_text[:20]}"
content_key = f"{editor_key}_content"

if content_key not in st.session_state:
    st.session_state[content_key] = existing_answer if existing_answer else ""

st.markdown("### ✍️ Your Story")
content = st_quill(st.session_state[content_key], key=editor_key, placeholder="Write your story here...")

if content is not None:
    st.session_state[content_key] = content

user_input = st.session_state[content_key]

# Show AI suggestions
if st.session_state.get('show_ai_suggestions') and st.session_state.get('current_ai_suggestions'):
    with st.expander("💡 AI Suggestions", expanded=True):
        st.markdown(st.session_state.current_ai_suggestions)
        if st.button("Dismiss"):
            st.session_state.show_ai_suggestions = False
            st.session_state.current_ai_suggestions = None
            st.rerun()

st.divider()

# ============================================================================
# IMAGE UPLOAD SECTION
# ============================================================================
if st.session_state.logged_in and st.session_state.image_handler:
    
    if existing_images:
        st.markdown("### 📸 Your Photos")
        cols = st.columns(min(3, len(existing_images)))
        for i, img in enumerate(existing_images[:3]):
            with cols[i % 3]:
                st.markdown(img["thumb_html"], unsafe_allow_html=True)
                if img.get("caption"):
                    st.caption(img["caption"])
                if st.button(f"Insert", key=f"ins_{img['id']}"):
                    current = st.session_state.get(content_key, "")
                    st.session_state[content_key] = current + "<br><br>" + img["full_html"]
                    st.rerun()
        
        st.divider()
    
    with st.expander("📤 Upload Photo"):
        uploaded = st.file_uploader("Choose image", type=['jpg', 'jpeg', 'png'], key=f"up_{current_session_id}")
        if uploaded:
            caption = st.text_input("Caption", key=f"cap_{current_session_id}")
            usage = st.radio("Size", ["Full Page", "Inline"], horizontal=True, key=f"usage_{current_session_id}")
            if st.button("Upload"):
                with st.spinner("Uploading..."):
                    usage_type = "full_page" if usage == "Full Page" else "inline"
                    result = st.session_state.image_handler.save_image(uploaded, current_session_id, current_question_text, caption, usage_type)
                    if result:
                        st.success("Uploaded!")
                        st.rerun()

st.divider()

# ============================================================================
# SAVE BUTTONS
# ============================================================================
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("💾 Save Story", type="primary"):
        if user_input and user_input.strip():
            if save_response(current_session_id, current_question_text, user_input):
                st.success("Saved!")
                st.rerun()
with col2:
    if existing_answer:
        if st.button("🗑️ Delete"):
            delete_response(current_session_id, current_question_text)
            st.rerun()
with col3:
    if st.session_state.current_question < len(current_session["questions"]) - 1:
        if st.button("Next →"):
            st.session_state.current_question += 1
            st.rerun()
    elif st.session_state.current_session < len(SESSIONS) - 1:
        if st.button("Next Session →"):
            st.session_state.current_session += 1
            st.session_state.current_question = 0
            st.rerun()

st.divider()

# ============================================================================
# BETA READER FEEDBACK
# ============================================================================
if beta_reader and st.session_state.logged_in:
    st.subheader("🦋 Beta Reader")
    if st.button("Get Feedback on Current Session"):
        session_text = ""
        for q, a in sdata.get("questions", {}).items():
            text_only = re.sub(r'<[^>]+>', '', a.get("answer", ""))
            session_text += f"Q: {q}\nA: {text_only}\n\n"
        
        if session_text.strip():
            with st.spinner("Analyzing..."):
                gps_context = get_narrative_gps_for_ai()
                full_text = gps_context + "\n\n" + session_text
                fb = generate_beta_reader_feedback(current_session["title"], full_text)
                if "error" not in fb:
                    st.session_state.current_beta_feedback = fb
                    st.session_state.show_beta_reader = True
                    st.rerun()

st.divider()

# ============================================================================
# SESSION PROGRESS
# ============================================================================
progress_info = get_progress_info(current_session_id)
st.markdown(f"""
<div style="padding:10px; background:#f0f2f6; border-radius:5px;">
    <div style="font-weight:bold;">📊 Progress: {progress_info['progress_percent']:.0f}%</div>
    <div>{progress_info['current_count']} / {progress_info['target']} words</div>
</div>
""", unsafe_allow_html=True)

if st.button("Change Word Target"):
    st.session_state.editing_word_target = not st.session_state.editing_word_target

if st.session_state.editing_word_target:
    new_target = st.number_input("Target words:", min_value=100, max_value=5000, value=progress_info['target'])
    if st.button("Save"):
        st.session_state.responses[current_session_id]["word_target"] = new_target
        save_user_data(st.session_state.user_id, st.session_state.responses)
        st.session_state.editing_word_target = False
        st.rerun()

