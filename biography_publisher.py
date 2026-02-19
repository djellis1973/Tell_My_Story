import streamlit as st
from datetime import datetime
import io
import base64
import os
import re
import html
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from fpdf import FPDF
import ebooklib
from ebooklib import epub
import tempfile

def clean_text(text):
    if not text:
        return text
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    text = text.replace('&#39;', "'")
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def show_celebration():
    st.balloons()
    st.success("🎉 Your book has been generated successfully!")

def generate_docx(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.first_line_indent = Inches(0.25)
    
    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title).font.size = Pt(42)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"by {author}").font.size = Pt(24)
    doc.add_page_break()
    
    # Copyright
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"© {datetime.now().year} {author}. All rights reserved.")
    doc.add_page_break()
    
    # TOC
    if include_toc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Table of Contents").font.size = Pt(18)
        sessions = {}
        for story in stories:
            s = story.get('session_title', 'Untitled')
            sessions[s] = True
        for s in sessions.keys():
            p = doc.add_paragraph()
            p.add_run(f"  {s}")
            p.paragraph_format.left_indent = Inches(0.5)
        doc.add_page_break()
    
    # Stories
    current = None
    for story in stories:
        session = story.get('session_title', 'Untitled')
        if session != current:
            current = session
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(session).font.size = Pt(16)
        
        if format_style == "interview" and story.get('question'):
            p = doc.add_paragraph()
            p.add_run(clean_text(story['question'])).bold = True
        
        if story.get('answer_text'):
            for para in clean_text(story['answer_text']).split('\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = Inches(0.25)
        
        doc.add_paragraph()
    
    bytes_io = io.BytesIO()
    doc.save(bytes_io)
    bytes_io.seek(0)
    return bytes_io.getvalue()

def generate_html(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(title)}</title>
    <style>
        body {{ font-family: Georgia; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ text-align: center; font-size: 42px; }}
        h2 {{ text-align: center; font-size: 28px; border-bottom: 1px solid #ccc; }}
        .author {{ text-align: center; font-style: italic; }}
        .question {{ font-weight: bold; margin-top: 30px; }}
    </style>
</head>
<body>
    <h1>{html.escape(title)}</h1>
    <p class="author">by {html.escape(author)}</p>
    <p style="text-align: center;">© {datetime.now().year}</p>
"""
    
    if include_toc:
        html_content += '<h2>Contents</h2><ul>'
        sessions = set(s.get('session_title', 'Untitled') for s in stories)
        for s in sessions:
            html_content += f'<li>{html.escape(s)}</li>'
        html_content += '</ul>'
    
    current = None
    for story in stories:
        session = story.get('session_title', 'Untitled')
        if session != current:
            current = session
            html_content += f'<h2>{html.escape(session)}</h2>'
        
        if format_style == "interview" and story.get('question'):
            html_content += f'<p class="question">{html.escape(clean_text(story["question"]))}</p>'
        
        if story.get('answer_text'):
            for para in clean_text(story['answer_text']).split('\n'):
                if para.strip():
                    html_content += f'<p>{html.escape(para.strip())}</p>'
        
        html_content += '<hr>'
    
    html_content += '</body></html>'
    return html_content

def generate_epub(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    book = epub.EpubBook()
    book.set_identifier('book123')
    book.set_title(title)
    book.set_language('en')
    book.add_author(author)
    
    chapters = []
    spine = ['nav']
    
    # Cover
    cover = epub.EpubHtml(title='Cover', file_name='cover.xhtml')
    cover.content = f'<h1>{title}</h1><h3>by {author}</h3>'
    book.add_item(cover)
    spine.append(cover)
    
    # Stories
    for i, story in enumerate(stories):
        session = story.get('session_title', f'Chapter {i+1}')
        content = f'<h1>{session}</h1>'
        if format_style == "interview" and story.get('question'):
            content += f'<p><b>{story["question"]}</b></p>'
        if story.get('answer_text'):
            content += f'<p>{story["answer_text"].replace(chr(10), "</p><p>")}</p>'
        
        chap = epub.EpubHtml(title=session, file_name=f'chap_{i}.xhtml')
        chap.content = content
        book.add_item(chap)
        spine.append(chap)
        chapters.append(chap)
    
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    bytes_io = io.BytesIO()
    epub.write_epub(bytes_io, book)
    bytes_io.seek(0)
    return bytes_io.getvalue()

class PDF(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font('Times', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font('Times', '', 12)
    pdf.set_auto_page_break(True, 25)
    
    # Cover
    pdf.set_font('Times', 'B', 42)
    pdf.cell(0, 100, title, 0, 1, 'C')
    pdf.set_font('Times', 'I', 24)
    pdf.cell(0, 20, f'by {author}', 0, 1, 'C')
    pdf.add_page()
    
    # Copyright
    pdf.set_font('Times', '', 12)
    pdf.cell(0, 10, f'© {datetime.now().year} {author}', 0, 1, 'C')
    
    # Stories
    for story in stories:
        session = story.get('session_title', 'Story')
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, session, 0, 1, 'C')
        pdf.ln(5)
        
        if format_style == "interview" and story.get('question'):
            pdf.set_font('Times', 'B', 12)
            pdf.multi_cell(0, 6, clean_text(story['question']))
            pdf.ln(3)
        
        if story.get('answer_text'):
            pdf.set_font('Times', '', 12)
            for para in clean_text(story['answer_text']).split('\n'):
                if para.strip():
                    pdf.multi_cell(0, 6, para.strip())
        
        pdf.ln(10)
    
    return pdf.output(dest='S').encode('latin-1')

def generate_rtf(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    rtf = "{\\rtf1\\ansi\\deff0 {\\fonttbl{\\f0 Times New Roman;}}\\f0\\fs24\n"
    rtf += f"\\pard\\qc \\b\\fs48 {title}\\b0\\par\n"
    rtf += f"\\pard\\qc \\i\\fs36 by {author}\\i0\\par\n\n"
    rtf += f"\\pard\\qc © {datetime.now().year} {author}\\par\n\\page\n"
    
    for story in stories:
        session = story.get('session_title', 'Story')
        rtf += f"\\pard\\qc \\b\\fs32 {session}\\b0\\par\n\n"
        if format_style == "interview" and story.get('question'):
            rtf += f"\\pard\\li720 \\b {story['question']}\\b0\\par\n"
        if story.get('answer_text'):
            for para in clean_text(story['answer_text']).split('\n'):
                if para.strip():
                    rtf += f"\\pard\\fi720 {para}\\par\n"
        rtf += "\\par\n"
    
    rtf += "}"
    return rtf.encode('utf-8')

def main():
    st.set_page_config(page_title="Biography Publisher", page_icon="📚", layout="wide")
    
    st.title("📚 Biography Publisher")
    st.markdown("Convert your stories into professional book formats")
    
    with st.sidebar:
        st.header("Book Settings")
        title = st.text_input("Book Title", "My Story Collection")
        author = st.text_input("Author Name", "Anonymous")
        
        st.subheader("Export Formats")
        export_docx = st.checkbox("📄 DOCX", value=True)
        export_html = st.checkbox("🌐 HTML", value=True)
        export_epub = st.checkbox("📱 EPUB", value=False)
        export_pdf = st.checkbox("📑 PDF", value=False)
        export_rtf = st.checkbox("📝 RTF", value=False)
        
        cover_choice = st.radio("Cover Type", ["simple", "uploaded"])
        cover_image = None
        if cover_choice == "uploaded":
            uploaded_cover = st.file_uploader("Upload Cover Image", type=['png', 'jpg', 'jpeg'])
            if uploaded_cover:
                cover_image = uploaded_cover.read()
        
        format_style = st.selectbox("Story Format", ["interview", "narrative"])
        include_toc = st.checkbox("Table of Contents", value=True)
        include_images = st.checkbox("Include Images", value=True)
    
    st.header("Add Your Stories")
    
    if 'stories' not in st.session_state:
        st.session_state.stories = []
    
    session_title = st.text_input("Session Title", "Session 1")
    if st.button("➕ New Session"):
        st.session_state.stories.append({
            'session_title': session_title,
            'question': '',
            'answer_text': '',
            'images': []
        })
        st.rerun()
    
    for i, story in enumerate(st.session_state.stories):
        with st.expander(f"📖 Story {i+1}: {story.get('session_title', 'Untitled')}"):
            story['session_title'] = st.text_input("Session Title", story['session_title'], key=f"s_{i}")
            story['question'] = st.text_area("Question", story.get('question', ''), key=f"q_{i}")
            story['answer_text'] = st.text_area("Answer", story.get('answer_text', ''), height=200, key=f"a_{i}")
            if st.button(f"Delete", key=f"del_{i}"):
                st.session_state.stories.pop(i)
                st.rerun()
    
    if st.button("➕ Add Another Story"):
        st.session_state.stories.append({
            'session_title': session_title,
            'question': '',
            'answer_text': '',
            'images': []
        })
        st.rerun()
    
    st.markdown("---")
    if st.button("📖 Generate Book", type="primary", use_container_width=True):
        if not st.session_state.stories:
            st.error("Add at least one story first.")
        else:
            with st.spinner("Generating..."):
                files = {}
                if export_docx:
                    files['docx'] = generate_docx(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                if export_html:
                    files['html'] = generate_html(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice).encode('utf-8')
                if export_epub:
                    files['epub'] = generate_epub(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                if export_pdf:
                    files['pdf'] = generate_pdf(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                if export_rtf:
                    files['rtf'] = generate_rtf(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                
                show_celebration()
                st.subheader("📥 Download Your Book")
                
                # SIMPLE - just show buttons one after another
                for fmt, data in files.items():
                    mime = {
                        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'html': 'text/html',
                        'epub': 'application/epub+zip',
                        'pdf': 'application/pdf',
                        'rtf': 'application/rtf'
                    }.get(fmt, 'application/octet-stream')
                    
                    st.download_button(
                        label=f"📥 Download {fmt.upper()}",
                        data=data,
                        file_name=f"{title.replace(' ', '_')}.{fmt}",
                        mime=mime,
                        use_container_width=True
                    )

if __name__ == "__main__":
    main()
